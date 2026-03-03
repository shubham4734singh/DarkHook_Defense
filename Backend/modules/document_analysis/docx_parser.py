# ================================================================
# docx_parser.py — DarkHOOK_ Defence
# Version  : 3.0 — Professional Grade
# Purpose  : Word document phishing detection using
#            14 industry-standard techniques
#
# Technique 1  -> File Type Validation
# Technique 2  -> Metadata Analysis
# Technique 3  -> Macro Detection
# Technique 4  -> Auto-Execution Detection
# Technique 5  -> VBA Behavior Analysis
# Technique 6  -> Macro Obfuscation Detection
# Technique 7  -> Embedded Object Analysis
# Technique 8  -> External Template Analysis
# Technique 9  -> Content Keyword Analysis
# Technique 10 -> URL and Hyperlink Analysis
# Technique 11 -> Attack Chain Inference
# Technique 12 -> Entropy and Payload Detection
# Technique 13 -> Reputation Matching
# Technique 14 -> Heuristic Risk Scoring
#
# Libraries: python-docx, oletools, zipfile,
#            hashlib, re, math
# ================================================================


# ----------------------------------------------------------------
# IMPORTS
# ----------------------------------------------------------------

import re
import os
import math
import hashlib
import zipfile
from urllib.parse import urlparse
from collections import Counter

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from oletools.olevba import VBA_Parser
    OLETOOLS_AVAILABLE = True
except ImportError:
    OLETOOLS_AVAILABLE = False


# ================================================================
# CONFIGURATION — Weights, keywords, patterns
# ================================================================

WEIGHTS = {
    # File structure
    "file_type_mismatch"         : 40,
    "corrupted_structure"        : 30,
    "double_extension"           : 35,
    "malformed_zip"              : 25,
    # Metadata
    "suspicious_metadata"        : 15,
    "wiped_metadata"             : 20,
    "metadata_mismatch"          : 20,
    "suspicious_template"        : 25,
    # Macro findings
    "malicious_macro"            : 40,
    "autoopen_macro"             : 35,
    "hidden_macro_stream"        : 35,
    # VBA behavior
    "suspicious_vba_api"         : 30,
    "powershell_in_vba"          : 40,
    "network_call_in_vba"        : 35,
    "file_system_access"         : 25,
    "registry_access"            : 30,
    "process_creation"           : 35,
    # Obfuscation
    "encoded_macro_payload"      : 35,
    "high_entropy_string"        : 25,
    "string_obfuscation"         : 25,
    "junk_code_detected"         : 15,
    # Embedded objects
    "embedded_ole_object"        : 30,
    "embedded_executable"        : 45,
    "embedded_script"            : 40,
    "double_extension_payload"   : 40,
    # External resources
    "external_template"          : 35,
    "external_image_tracker"     : 20,
    "suspicious_relationship"    : 25,
    "hidden_relationship"        : 30,
    # Content
    "phishing_keyword"           : 10,
    "urgent_tone_detected"       : 15,
    "financial_terms_detected"   : 15,
    "credential_harvesting"      : 20,
    "enable_macro_lure"          : 35,
    "repeated_cta"               : 15,
    # URL findings
    "suspicious_url"             : 15,
    "ip_based_url"               : 30,
    "shortened_url"              : 20,
    "suspicious_tld"             : 20,
    "at_symbol_trick"            : 25,
    "hidden_hyperlink"           : 25,
    "mismatched_anchor"          : 25,
    # Attack chain
    "dropper_pattern"            : 40,
    "download_execute_pattern"   : 40,
    "multistage_indicator"       : 35,
    # Reputation
    "known_malicious_hash"       : 100,
    "known_macro_signature"      : 45,
}


# ----------------------------------------------------------------
# PHISHING KEYWORDS — Technique 9
# ----------------------------------------------------------------

PHISHING_KEYWORDS = {

    "account_threats": [
        "verify your account",
        "confirm your account",
        "account suspended",
        "account will be closed",
        "account has been compromised",
        "account limited",
        "reactivate your account",
        "unlock your account",
        "account will be terminated",
        "unusual activity detected",
    ],

    "urgency_phrases": [
        "urgent action required",
        "immediate action required",
        "act now",
        "respond immediately",
        "expires today",
        "final warning",
        "last chance",
        "do not ignore",
        "failure to respond",
        "within 24 hours",
        "time sensitive",
        "don't delay",
    ],

    "credential_harvesting": [
        "enter your password",
        "confirm your password",
        "reset your password",
        "password expired",
        "enter your credentials",
        "verify your identity",
        "sign in to continue",
        "login credentials required",
        "enter your username",
        "enter your email",
    ],

    "financial_terms": [
        "bank account details",
        "credit card details",
        "debit card number",
        "billing information required",
        "payment details required",
        "update payment method",
        "transaction failed",
        "wire transfer",
        "gift card",
        "bitcoin payment",
        "refund pending",
    ],

    "reward_tricks": [
        "you have won",
        "prize money",
        "claim your reward",
        "lottery winner",
        "congratulations you won",
        "cash prize",
        "free gift",
    ],

    "legal_threats": [
        "legal action will be taken",
        "police complaint filed",
        "court notice",
        "government notice",
        "income tax department",
        "irs notice",
        "tax refund",
        "warrant issued",
        "arrest warrant",
    ],

    "india_specific": [
        "aadhar number",
        "pan card details",
        "kyc verification",
        "kyc update required",
        "upi details",
        "enter otp",
        "otp verification",
    ],

    "fake_security_alerts": [
        "your computer is infected",
        "virus detected",
        "security breach",
        "unauthorized access",
        "install this update",
        "download the security patch",
        "suspicious login attempt",
    ],

    "enable_macro_lures": [
        "enable macros to view",
        "enable editing to continue",
        "enable content to view",
        "click enable to continue",
        "must enable macros",
        "enable to view document",
        "protected document",
        "click enable content",
    ],

    "download_tricks": [
        "click the link below",
        "click here to verify",
        "download the attachment",
        "open the document",
        "view your invoice",
        "access your account here",
        "download now",
        "click here immediately",
    ],
}


# ----------------------------------------------------------------
# DANGEROUS VBA APIs — Technique 5
# ----------------------------------------------------------------

DANGEROUS_VBA_APIS = {

    "shell_execution": [
        "shell(",
        "shell ",
        "wscript.shell",
        "shell.application",
        "shellexecute",
        "winexec(",
        'createobject("wscript',
    ],

    "powershell": [
        "powershell",
        "powershell.exe",
        "-encodedcommand",
        "-enc ",
        "invoke-expression",
        "iex(",
        "invoke-webrequest",
    ],

    "network_calls": [
        "xmlhttp",
        "xmlhttprequest",
        "urldownloadtofile",
        "winhttprequest",
        "internetexplorer.application",
        "msxml2.xmlhttp",
        "http.open",
    ],

    "file_system": [
        "filesystemobject",
        "createtextfile",
        "opentextfile",
        "copyfile",
        "deletefile",
        "movefile",
        "kill(",
        "filecopy",
    ],

    "registry": [
        "regwrite",
        "regread",
        "regdelete",
        "hkey_local_machine",
        "hkey_current_user",
        "registry",
        "regedit",
    ],

    "process_creation": [
        "createprocess",
        "shellexecute",
        "getobject(",
        "wmi",
        "win32_process",
    ],
}


# AutoOpen names — Technique 4
AUTOOPEN_NAMES = [
    "autoopen",
    "auto_open",
    "document_open",
    "documentopen",
    "autoclose",
    "auto_close",
    "document_close",
    "autoexec",
    "workbook_open",
    "auto_exec",
    "worksheet_activate",
]


# Known malicious macro signatures — Technique 13
KNOWN_MACRO_SIGNATURES = [
    "powershell -nop -w hidden",
    "iex(new-object net.webclient",
    "urldownloadtofile",
    'shell("cmd /c',
    'wscript.shell").run',
    'createobject("adodb.stream")',
    'environ("appdata")',
]


# URL shorteners — Technique 10
URL_SHORTENERS = [
    "bit.ly",
    "tinyurl.com",
    "t.co",
    "goo.gl",
    "ow.ly",
    "is.gd",
    "buff.ly",
    "rebrand.ly",
    "cutt.ly",
    "shorturl.at",
    "tiny.cc",
    "rb.gy",
]


# Suspicious TLDs — Technique 10
SUSPICIOUS_TLDS = [
    ".xyz",
    ".top",
    ".ru",
    ".tk",
    ".ml",
    ".ga",
    ".cf",
    ".gq",
    ".pw",
    ".click",
    ".download",
    ".loan",
    ".work",
    ".party",
]


# ================================================================
# HELPER FUNCTIONS
# ================================================================

def calculate_entropy(text):
    """
    Shannon entropy measures randomness.
    Normal text  -> 3.0 to 4.5
    Encoded data -> 6.0 to 8.0
    """
    if not text or len(text) < 20:
        return 0.0
    counter = Counter(text)
    length  = len(text)
    entropy = -sum(
        (c / length) * math.log2(c / length)
        for c in counter.values()
    )
    return round(entropy, 2)


def is_ip_url(url):
    """Returns True if URL uses IP instead of domain"""
    try:
        host = urlparse(url).netloc
        if ":" in host:
            host = host.split(":")[0]
        return bool(re.match(
            r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$', host
        ))
    except Exception:
        return False


def analyze_url(url):
    """Full URL analysis — returns findings and details"""
    url_findings = []
    url_details  = []

    try:
        parsed = urlparse(url)
        domain = parsed.netloc.lower()
        path   = parsed.path.lower()

        # Check 1 — IP URL
        if is_ip_url(url):
            url_findings.append("ip_based_url")
            url_details.append("IP-based URL: " + url)

        # Check 2 — URL shortener
        for s in URL_SHORTENERS:
            if s in domain:
                url_findings.append("shortened_url")
                url_details.append("URL shortener (" + s + "): " + url)
                break

        # Check 3 — Suspicious TLD
        for tld in SUSPICIOUS_TLDS:
            if domain.endswith(tld):
                url_findings.append("suspicious_tld")
                url_details.append("Suspicious TLD (" + tld + "): " + url)
                break

        # Check 4 — @ trick
        if "@" in url:
            url_findings.append("at_symbol_trick")
            url_details.append("@ redirect trick: " + url)

        # Check 5 — HTTP insecure
        if url.startswith("http://"):
            url_findings.append("suspicious_url")
            url_details.append("Insecure HTTP: " + url)

        # Check 6 — Long URL
        if len(url) > 200:
            url_findings.append("suspicious_url")
            url_details.append("Suspicious long URL: " + url[:60] + "...")

        # Check 7 — Fake login keywords in path
        fake_words = [
            "login", "signin", "verify", "secure",
            "account", "update", "confirm", "banking",
            "password", "credential",
        ]
        for word in fake_words:
            if word in path:
                url_findings.append("suspicious_url")
                url_details.append(
                    "Login keyword in URL (" + word + "): " + url
                )
                break

        # Check 8 — Homograph domain
        if re.search(r'[àáâãäåæçèéêëìíîïðñòóôõöøùúûü]', domain):
            url_findings.append("suspicious_url")
            url_details.append("Homograph domain: " + domain)

        # Check 9 — Double HTTP
        if url.count("http") > 1:
            url_findings.append("suspicious_url")
            url_details.append("Double HTTP obfuscation: " + url)

    except Exception as e:
        url_details.append("URL error: " + str(e))

    return url_findings, url_details


# ================================================================
# TECHNIQUE 1 — File Type and Structure Validation
# ================================================================

def technique1_file_validation(file_path):
    """
    Verifies file is actually a real Word document.
    Attackers rename malicious files to .docx to trick users.
    """
    findings = []
    details  = []

    details.append("--- TECHNIQUE 1: FILE VALIDATION ---")

    filename = os.path.basename(file_path)
    ext      = os.path.splitext(filename)[1].lower()

    # Check 1 — Valid extension
    valid_extensions = [".docx", ".docm", ".doc", ".dotx", ".dotm"]
    if ext not in valid_extensions:
        findings.append("file_type_mismatch")
        details.append("Invalid extension: " + ext)
    else:
        details.append("Extension valid: " + ext)

    # Check 2 — .docm always has macros
    if ext == ".docm":
        findings.append("malicious_macro")
        details.append("CRITICAL: .docm = macro-enabled document!")

    # Check 3 — Double extension trick
    # We ignore version numbers like 2.0, 1.5, 3.2
    # Real double extension looks like: invoice.pdf.docx
    name_without_ext = os.path.splitext(filename)[0]
    if "." in name_without_ext:

        # Get the part after the last dot in name
        # Example: "invoice.pdf" -> suspicious part = "pdf"
        # Example: "policy 2.0"  -> suspicious part = "0" (number = safe)
        part_after_dot = name_without_ext.split(".")[-1].strip()

        # Check if it is a version number (all digits like "0", "1", "5")
        is_version_number = part_after_dot.isdigit()

        # Check if it looks like "2.0" pattern (digit.digit)
        is_version_pattern = bool(
            re.match(r'^\d+$', part_after_dot)
        )

        # Check if it is a known dangerous extension hiding inside name
        dangerous_inner_exts = [
            "pdf", "exe", "dll", "bat", "cmd",
            "ps1", "vbs", "js", "hta", "scr",
            "zip", "rar", "7z", "iso",
        ]
        is_dangerous_ext = part_after_dot.lower() in dangerous_inner_exts

        if is_dangerous_ext:
            # Real double extension attack
            findings.append("double_extension")
            details.append(
                "Double extension detected: " + filename +
                " — dangerous extension hiding inside name!"
            )
        elif is_version_number or is_version_pattern:
            # Just a version number — safe
            details.append(
                "Note: Dot in filename is version number (" +
                part_after_dot + ") — not a double extension"
            )
        else:
            # Unknown pattern — warn but lower severity
            details.append(
                "Note: Dot found in filename — " +
                "verify manually: " + filename
            )

    # Check 4 — Verify actual file is a valid ZIP
    if ext in [".docx", ".docm", ".dotx", ".dotm"]:
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                names = z.namelist()
                if "word/document.xml" not in names:
                    findings.append("corrupted_structure")
                    details.append(
                        "Malformed DOCX: word/document.xml missing"
                    )
                else:
                    details.append("DOCX ZIP structure valid")
        except zipfile.BadZipFile:
            findings.append("malformed_zip")
            details.append(
                "File is NOT a valid ZIP/DOCX — possible disguised malware!"
            )
        except Exception as e:
            details.append("ZIP check error: " + str(e))

    details.append("Technique 1 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 2 — Metadata Analysis
# ================================================================

def technique2_metadata(file_path):
    """
    Analyses hidden metadata inside Word document.
    Suspicious metadata = possible forged or phishing document.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 2: METADATA ANALYSIS ---")

    if not DOCX_AVAILABLE:
        details.append("python-docx not available")
        return findings, details

    try:
        doc  = Document(file_path)
        core = doc.core_properties

        author   = core.author   or ""
        created  = core.created
        modified = core.modified
        revision = core.revision or 0
        template = ""

        # Try to get template from relationships
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/_rels/settings.xml.rels" in z.namelist():
                    with z.open("word/_rels/settings.xml.rels") as f:
                        template = f.read().decode("utf-8", errors="ignore")
        except Exception:
            pass

        details.append("Author   : " + (author or "EMPTY"))
        details.append("Created  : " + str(created))
        details.append("Modified : " + str(modified))
        details.append("Revision : " + str(revision))

        # Check 1 — Empty or wiped author
        if not author or author.strip() == "":
            findings.append("wiped_metadata")
            details.append(
                "Author field empty — metadata may have been wiped deliberately"
            )

        # Check 2 — Generic suspicious author names
        suspicious_names = [
            "admin", "user", "test", "administrator",
            "temp", "unknown", "abc", "owner",
        ]
        if author.lower().strip() in suspicious_names:
            findings.append("suspicious_metadata")
            details.append("Suspicious author name: " + author)

        # Check 3 — Created and modified same time
        if created and modified:
            diff = abs((modified - created).total_seconds())
            if diff < 2:
                findings.append("suspicious_metadata")
                details.append(
                    "Created and modified at same time — "
                    "may be auto-generated phishing document"
                )

        # Check 4 — Revision count = 1
        if revision == 1:
            findings.append("suspicious_metadata")
            details.append(
                "Revision count = 1 — "
                "document may have been generated not typed"
            )

        # Check 5 — External template
        if "http://" in template or "https://" in template:
            findings.append("external_template")
            details.append(
                "External template URL in settings — "
                "loads malicious template from internet on open!"
            )
            url_match = re.search(r'Target="(https?://[^"]+)"', template)
            if url_match:
                details.append("External template URL: " + url_match.group(1))

    except Exception as e:
        details.append("Metadata error: " + str(e))

    details.append("Technique 2 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUES 3, 4, 5, 6 — Macro Analysis
# ================================================================

def techniques3456_macro_analysis(file_path):
    """
    Technique 3 -> Macro presence detection
    Technique 4 -> Auto-execution detection
    Technique 5 -> VBA behavior analysis
    Technique 6 -> Macro obfuscation detection
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUES 3-6: MACRO ANALYSIS ---")

    if not OLETOOLS_AVAILABLE:
        details.append("oletools not available — skipping macro analysis")
        return findings, details

    try:
        vba = VBA_Parser(file_path)

        # -----------------------------------------------
        # TECHNIQUE 3 — Macro presence
        # -----------------------------------------------

        if not vba.detect_vba_macros():
            details.append("Technique 3: No VBA macros detected")
            vba.close()
            return findings, details

        findings.append("malicious_macro")
        details.append("Technique 3: VBA macros DETECTED!")

        # Check for hidden macro streams
        try:
            results = vba.analyze_macros()
            for kw_type, kw_keyword, kw_description in results:
                if "suspicious" in kw_type.lower():
                    findings.append("hidden_macro_stream")
                    details.append(
                        "Hidden/suspicious stream: " + str(kw_description)
                    )
        except Exception:
            pass

        # Extract all macro code
        all_vba_code = ""
        for (filename, stream_path,
             vba_filename, vba_code) in vba.extract_macros():
            all_vba_code += vba_code.lower() + "\n"

        # -----------------------------------------------
        # TECHNIQUE 4 — Auto-execution detection
        # -----------------------------------------------

        for auto_name in AUTOOPEN_NAMES:
            if auto_name in all_vba_code:
                findings.append("autoopen_macro")
                details.append(
                    "Technique 4: AutoOpen macro: '" +
                    auto_name + "' — RUNS ON DOCUMENT OPEN!"
                )

        # -----------------------------------------------
        # TECHNIQUE 5 — VBA behavior analysis
        # -----------------------------------------------

        for category, apis in DANGEROUS_VBA_APIS.items():
            for api in apis:
                if api in all_vba_code:
                    if category == "shell_execution":
                        findings.append("suspicious_vba_api")
                        details.append("Technique 5 [Shell]: " + api)
                    elif category == "powershell":
                        findings.append("powershell_in_vba")
                        details.append("Technique 5 [PowerShell]: " + api)
                    elif category == "network_calls":
                        findings.append("network_call_in_vba")
                        details.append("Technique 5 [Network]: " + api)
                    elif category == "file_system":
                        findings.append("file_system_access")
                        details.append("Technique 5 [FileSystem]: " + api)
                    elif category == "registry":
                        findings.append("registry_access")
                        details.append("Technique 5 [Registry]: " + api)
                    elif category == "process_creation":
                        findings.append("process_creation")
                        details.append("Technique 5 [Process]: " + api)

        # -----------------------------------------------
        # TECHNIQUE 6 — Obfuscation detection
        # -----------------------------------------------

        # Check 1 — Base64 encoded strings
        b64_pattern = re.compile(r'[A-Za-z0-9+/]{50,}={0,2}')
        b64_matches = b64_pattern.findall(all_vba_code)
        if b64_matches:
            findings.append("encoded_macro_payload")
            details.append(
                "Technique 6 [Base64]: " + b64_matches[0][:50] + "..."
            )

        # Check 2 — Hex encoded strings
        hex_pattern = re.compile(r'[0-9a-f]{40,}')
        hex_matches = hex_pattern.findall(all_vba_code)
        if hex_matches:
            findings.append("encoded_macro_payload")
            details.append(
                "Technique 6 [Hex]: " + hex_matches[0][:50] + "..."
            )

        # Check 3 — String concatenation obfuscation
        concat_patterns = [
            r'"\s*&\s*"',
            r'chr\(\d+\)',
            r'chrw\(\d+\)',
        ]
        for pattern in concat_patterns:
            if re.search(pattern, all_vba_code):
                findings.append("string_obfuscation")
                details.append(
                    "Technique 6 [Obfuscation]: String pattern: " + pattern
                )

        # Check 4 — High entropy
        entropy = calculate_entropy(all_vba_code[:1000])
        if entropy > 5.5:
            findings.append("high_entropy_string")
            details.append(
                "Technique 6 [Entropy]: High entropy: " +
                str(entropy) + " — obfuscated"
            )

        # Check 5 — Known malicious signatures
        for signature in KNOWN_MACRO_SIGNATURES:
            if signature in all_vba_code:
                findings.append("known_macro_signature")
                details.append(
                    "Technique 13 [Known Signature]: " + signature[:50]
                )

        vba.close()

    except Exception as e:
        details.append("Macro analysis error: " + str(e))

    details.append("Techniques 3-6 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 7 — Embedded Object and OLE Analysis
# ================================================================

def technique7_embedded_objects(file_path):
    """
    Checks for hidden files embedded inside Word document.
    Embedded executables are extremely dangerous.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 7: EMBEDDED OBJECT ANALYSIS ---")

    try:
        with zipfile.ZipFile(file_path, "r") as z:
            all_files = z.namelist()

            # Check for OLE embedded objects
            ole_files = [
                f for f in all_files
                if "embeddings" in f.lower() or "oleobject" in f.lower()
            ]

            if ole_files:
                for ole_file in ole_files:
                    findings.append("embedded_ole_object")
                    details.append("Embedded OLE object: " + ole_file)

                    # Check for dangerous extensions
                    dangerous_exts = [
                        ".exe", ".dll", ".bat", ".cmd",
                        ".ps1", ".vbs", ".js", ".hta",
                        ".scr", ".pif", ".com",
                    ]
                    for ext in dangerous_exts:
                        if ole_file.lower().endswith(ext):
                            findings.append("embedded_executable")
                            details.append(
                                "CRITICAL: Embedded executable: " + ole_file
                            )

                    # Check double extension in embedded files
                    base = os.path.splitext(ole_file)[0]
                    if "." in os.path.basename(base):
                        findings.append("double_extension_payload")
                        details.append(
                            "Double extension in embedded object: " + ole_file
                        )

            # Check embedded media for scripts
            script_files = [
                f for f in all_files
                if any(
                    f.lower().endswith(ext)
                    for ext in [".js", ".vbs", ".ps1", ".bat", ".cmd", ".hta"]
                )
            ]

            for sf in script_files:
                findings.append("embedded_script")
                details.append("CRITICAL: Embedded script: " + sf)

            # Check for actual executables (MZ header)
            for fname in all_files:
                try:
                    data = z.read(fname)
                    if data[:2] == b"MZ":
                        findings.append("embedded_executable")
                        details.append(
                            "CRITICAL: EXE signature (MZ) found in: " + fname
                        )
                except Exception:
                    pass

            if not ole_files and not script_files:
                details.append("No suspicious embedded objects")

    except Exception as e:
        details.append("Embedded object error: " + str(e))

    details.append("Technique 7 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 8 — External Template and Resource Analysis
# ================================================================

def technique8_external_resources(file_path):
    """
    Checks if document loads external resources.
    Attackers host malicious templates/payloads online.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 8: EXTERNAL RESOURCE ANALYSIS ---")

    try:
        with zipfile.ZipFile(file_path, "r") as z:

            rel_files = [f for f in z.namelist() if f.endswith(".rels")]

            for rel_file in rel_files:
                with z.open(rel_file) as f:
                    content = f.read().decode("utf-8", errors="ignore")

                    # External template
                    if "attachedTemplate" in content:
                        url_match = re.search(
                            r'Target="(https?://[^"]+)"', content
                        )
                        if url_match:
                            findings.append("external_template")
                            details.append(
                                "CRITICAL: External template loads from: " +
                                url_match.group(1)
                            )

                    # External image trackers
                    if "image" in rel_file.lower():
                        ext_urls = re.findall(
                            r'Target="(https?://[^"]+)"', content
                        )
                        for url in ext_urls:
                            findings.append("external_image_tracker")
                            details.append(
                                "External image URL (possible tracker): " + url
                            )

                    # All external relationships
                    all_ext = re.findall(
                        r'TargetMode="External"[^>]*Target="([^"]+)"',
                        content
                    )
                    for target in all_ext:
                        if target.startswith("http"):
                            findings.append("suspicious_relationship")
                            details.append(
                                "External relationship: " + target
                            )

            # Check settings.xml for external references
            if "word/settings.xml" in z.namelist():
                with z.open("word/settings.xml") as f:
                    settings = f.read().decode("utf-8", errors="ignore")

                    if "http://" in settings or "https://" in settings:

                        # Extract all URLs from settings.xml
                        urls_in_settings = re.findall(
                            r'https?://[^\s"<>]+', settings
                        )

                        # These are Microsoft/Office safe domains
                        # They appear in ALL normal Word documents
                        safe_domains = [
                            "schemas.openxmlformats.org",
                            "schemas.microsoft.com",
                            "purl.org",
                            "www.w3.org",
                            "schemas.openformats.org",
                            "microsoft.com",
                            "office.com",
                            "officeapps.live.com",
                        ]

                        for url in urls_in_settings:
                            url_lower = url.lower()

                            # Check if it is a known safe Microsoft URL
                            is_safe = any(
                                safe in url_lower
                                for safe in safe_domains
                            )

                            if not is_safe:
                                # Real suspicious URL found!
                                findings.append("hidden_relationship")
                                details.append(
                                    "Suspicious external URL in "
                                    "settings.xml: " + url
                                )

    except Exception as e:
        details.append("External resource error: " + str(e))

    if not findings:
        details.append("No external resources detected")

    details.append("Technique 8 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUES 9 and 10 — Content and URL Analysis
# ================================================================

def techniques910_content_url(file_path):
    """
    Technique 9  -> Phishing keyword analysis
    Technique 10 -> URL and hyperlink analysis
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUES 9-10: CONTENT AND URL ANALYSIS ---")

    if not DOCX_AVAILABLE:
        details.append("python-docx not available")
        return findings, details

    try:
        doc       = Document(file_path)
        full_text = ""
        all_urls  = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text += para.text.lower() + " "

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text.lower() + " "

        # Extract hyperlinks
        try:
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel.target_ref
                    if url.startswith("http"):
                        all_urls.append(("hyperlink", url, ""))
        except Exception:
            pass

        # Find URLs in plain text
        url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+')
        for url in url_pattern.findall(full_text):
            all_urls.append(("text", url, ""))

        # --------------------------------------------------
        # TECHNIQUE 10 — URL Analysis
        # --------------------------------------------------

        details.append("Total URLs found: " + str(len(all_urls)))

        if len(all_urls) > 10:
            findings.append("suspicious_url")
            details.append(
                "High URL count: " + str(len(all_urls)) + " URLs — suspicious"
            )

        for url_type, url, visible_text in all_urls:
            url_findings, url_details = analyze_url(url)
            if url_findings:
                findings.extend(url_findings)
                for d in url_details:
                    details.append("  [" + url_type + "] " + d)
            else:
                details.append("  [" + url_type + "] Link: " + url)

            if visible_text and url:
                vis_lower = visible_text.lower()
                url_lower = url.lower()
                if "http" in vis_lower and vis_lower.strip() != url_lower.strip():
                    findings.append("mismatched_anchor")
                    details.append(
                        "Mismatched anchor: shows '" +
                        visible_text[:40] + "' links to '" + url[:40] + "'"
                    )

        # --------------------------------------------------
        # TECHNIQUE 9 — Keyword Analysis
        # --------------------------------------------------

        urgency_count    = 0
        financial_count  = 0
        credential_count = 0
        enable_count     = 0
        cta_count        = 0

        for category, keywords in PHISHING_KEYWORDS.items():
            for keyword in keywords:
                count = full_text.count(keyword)
                if count > 0:
                    findings.append("phishing_keyword")
                    details.append(
                        "[" + category + "] '" + keyword + "' x" + str(count)
                    )
                    if category == "urgency_phrases":
                        urgency_count += count
                    elif category == "financial_terms":
                        financial_count += count
                    elif category == "credential_harvesting":
                        credential_count += count
                    elif category == "enable_macro_lures":
                        enable_count += count
                    elif category == "download_tricks":
                        cta_count += count

        if urgency_count >= 2:
            findings.append("urgent_tone_detected")
            details.append("Urgency tone: " + str(urgency_count) + " phrases")

        if financial_count >= 2:
            findings.append("financial_terms_detected")
            details.append(
                "Financial targeting: " + str(financial_count) + " terms"
            )

        if credential_count >= 1:
            findings.append("credential_harvesting")
            details.append(
                "Credential harvesting: " + str(credential_count) + " phrases"
            )

        if enable_count >= 1:
            findings.append("enable_macro_lure")
            details.append(
                "CRITICAL: Enable macro lure: " + str(enable_count) +
                "x — tricks user into enabling macros!"
            )

        if cta_count >= 3:
            findings.append("repeated_cta")
            details.append("Repeated call-to-action: " + str(cta_count) + "x")

    except Exception as e:
        details.append("Content/URL error: " + str(e))

    details.append("Techniques 9-10 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 11 — Attack Chain Inference
# ================================================================

def technique11_attack_chain(all_findings):
    """
    Infers the possible attack sequence from findings.
    Helps security teams understand what happens next.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 11: ATTACK CHAIN INFERENCE ---")

    finding_set = set(all_findings)

    # Pattern 1 — Macro download-execute chain
    if "malicious_macro" in finding_set and "network_call_in_vba" in finding_set:
        findings.append("download_execute_pattern")
        details.append(
            "ATTACK CHAIN: Macro -> Network call -> Download payload -> Execute"
        )

    # Pattern 2 — Template injection chain
    if "external_template" in finding_set:
        findings.append("multistage_indicator")
        details.append(
            "ATTACK CHAIN: Template injection -> Load remote macro -> Execute on open"
        )

    # Pattern 3 — Social engineering + macro
    if "enable_macro_lure" in finding_set and "malicious_macro" in finding_set:
        findings.append("dropper_pattern")
        details.append(
            "ATTACK CHAIN: Show blurred content -> User enables macros -> Payload executes"
        )

    # Pattern 4 — Credential theft
    if "credential_harvesting" in finding_set and (
        "suspicious_url" in finding_set or "ip_based_url" in finding_set
    ):
        findings.append("dropper_pattern")
        details.append(
            "ATTACK CHAIN: Credential lure -> Redirect to fake login -> Steal credentials"
        )

    # Pattern 5 — Embedded dropper
    if "embedded_executable" in finding_set or "embedded_ole_object" in finding_set:
        findings.append("dropper_pattern")
        details.append(
            "ATTACK CHAIN: Embedded payload -> User opens document -> Execute dropped file"
        )

    if not findings:
        details.append("No clear attack chain identified")

    details.append("Technique 11 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 12 — Entropy and Encrypted Payload Detection
# ================================================================

def technique12_entropy(file_path):
    """
    Checks for encrypted, compressed or obfuscated
    data sections that hide malicious payloads.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 12: ENTROPY ANALYSIS ---")

    try:
        raw_content = ""

        with zipfile.ZipFile(file_path, "r") as z:
            for fname in z.namelist():
                try:
                    with z.open(fname) as f:
                        data = f.read()

                        # Check binary files for high entropy
                        if not fname.endswith(".xml"):
                            entropy = calculate_entropy(
                                data[:500].decode("latin-1", errors="ignore")
                            )
                            if entropy > 7.0:
                                findings.append("high_entropy_string")
                                details.append(
                                    "High entropy in " + fname + ": " +
                                    str(entropy) + " — possible encrypted payload"
                                )

                        # Collect XML content
                        if fname.endswith(".xml"):
                            raw_content += data.decode("utf-8", errors="ignore")

                except Exception:
                    pass

        # Check for suspicious Base64 blobs
        b64_pattern = re.compile(r'[A-Za-z0-9+/]{100,}={0,2}')
        b64_matches = b64_pattern.findall(raw_content)
        suspicious_b64 = [
            m for m in b64_matches
            if calculate_entropy(m) > 5.5
        ]

        if suspicious_b64:
            findings.append("encoded_macro_payload")
            details.append(
                "Suspicious Base64 blob (" + str(len(suspicious_b64)) +
                " found): " + suspicious_b64[0][:50] + "..."
            )

        # Check for hex encoded content
        hex_pattern = re.compile(r'[0-9a-fA-F]{80,}')
        hex_matches = hex_pattern.findall(raw_content)
        if hex_matches:
            findings.append("encoded_macro_payload")
            details.append(
                "Hex encoded content: " + hex_matches[0][:50] + "..."
            )

        if not findings:
            details.append("No suspicious encrypted payloads")

    except Exception as e:
        details.append("Entropy error: " + str(e))

    details.append("Technique 12 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 13 — Reputation and Known Threat Matching
# ================================================================

def technique13_reputation(file_path, all_vba_code=""):
    """
    Compares file and macro code against known
    malicious patterns and signatures.
    """
    findings = []
    details  = []

    details.append("")
    details.append("--- TECHNIQUE 13: REPUTATION MATCHING ---")

    try:
        with open(file_path, "rb") as f:
            sha256 = hashlib.sha256(f.read()).hexdigest()

        details.append("SHA256: " + sha256)
        details.append("Check hash at: https://www.virustotal.com")

        if all_vba_code:
            vba_lower = all_vba_code.lower()
            for sig in KNOWN_MACRO_SIGNATURES:
                if sig in vba_lower:
                    findings.append("known_macro_signature")
                    details.append("Known malicious signature: " + sig[:50])

        if not findings:
            details.append("No known malicious signatures matched")

    except Exception as e:
        details.append("Reputation error: " + str(e))

    details.append("Technique 13 findings: " + str(len(findings)))
    return findings, details


# ================================================================
# TECHNIQUE 14 — Heuristic Risk Scoring
# ================================================================

def technique14_scoring(all_findings):
    """
    Converts all findings into a final weighted score.
    Exactly like how real antivirus engines work.
    """
    total_score = 0
    breakdown   = {}

    for finding in all_findings:
        weight = WEIGHTS.get(finding, 5)
        total_score += weight

        if finding in breakdown:
            breakdown[finding]["count"] += 1
            breakdown[finding]["score"] += weight
        else:
            breakdown[finding] = {
                "count" : 1,
                "score" : weight
            }

    # Cap at 100
    total_score = min(total_score, 100)

    if total_score < 30:
        verdict  = "Low Risk"
        severity = "LOW"
    elif total_score < 60:
        verdict  = "Medium Risk"
        severity = "MEDIUM"
    elif total_score < 80:
        verdict  = "High Risk"
        severity = "HIGH"
    else:
        verdict  = "Critical — Likely Phishing"
        severity = "CRITICAL"

    return total_score, verdict, severity, breakdown


# ================================================================
# MAIN FUNCTION — parse_docx
# Called by app.py when user uploads Word file
# ================================================================

def parse_docx(file_path):
    """
    file_path = full path to uploaded Word file
    Returns   = complete analysis result dict
    """

    all_findings = []
    all_details  = []
    sha256_hash  = ""

    try:

        # Header
        all_details.append("=" * 55)
        all_details.append("DARKHOOK_ DEFENCE — WORD DOCUMENT ANALYSIS")
        all_details.append("=" * 55)
        all_details.append("File: " + os.path.basename(file_path))

        # SHA256
        with open(file_path, "rb") as f:
            sha256_hash = hashlib.sha256(f.read()).hexdigest()

        all_details.append("SHA256: " + sha256_hash)
        all_details.append("")

        # -----------------------------------------------
        # Run all 14 techniques
        # -----------------------------------------------

        # Technique 1
        f1, d1 = technique1_file_validation(file_path)
        all_findings.extend(f1)
        all_details.extend(d1)

        # Technique 2
        f2, d2 = technique2_metadata(file_path)
        all_findings.extend(f2)
        all_details.extend(d2)

        # Techniques 3-6
        f3, d3 = techniques3456_macro_analysis(file_path)
        all_findings.extend(f3)
        all_details.extend(d3)

        # Technique 7
        f7, d7 = technique7_embedded_objects(file_path)
        all_findings.extend(f7)
        all_details.extend(d7)

        # Technique 8
        f8, d8 = technique8_external_resources(file_path)
        all_findings.extend(f8)
        all_details.extend(d8)

        # Techniques 9-10
        f9, d9 = techniques910_content_url(file_path)
        all_findings.extend(f9)
        all_details.extend(d9)

        # Technique 11 — needs all findings so far
        f11, d11 = technique11_attack_chain(all_findings)
        all_findings.extend(f11)
        all_details.extend(d11)

        # Technique 12
        f12, d12 = technique12_entropy(file_path)
        all_findings.extend(f12)
        all_details.extend(d12)

        # Technique 13
        f13, d13 = technique13_reputation(file_path)
        all_findings.extend(f13)
        all_details.extend(d13)

        # Technique 14 — Final scoring
        score, verdict, severity, breakdown = technique14_scoring(all_findings)

        all_details.append("")
        all_details.append("--- TECHNIQUE 14: HEURISTIC SCORING ---")
        all_details.append("Total techniques run : 14")
        all_details.append("Total findings       : " + str(len(all_findings)))
        all_details.append("Danger score         : " + str(score) + "/100")
        all_details.append("Severity             : " + severity)
        all_details.append("Verdict              : " + verdict)
        all_details.append("")
        all_details.append("Score breakdown:")

        for finding, data in breakdown.items():
            all_details.append(
                "  " + finding.ljust(35) +
                " count=" + str(data["count"]) +
                " score=" + str(data["score"])
            )

    except Exception as error:
        all_details.append("Critical error: " + str(error))

    return {
        "findings" : all_findings,
        "details"  : all_details,
        "sha256"   : sha256_hash,
    }
